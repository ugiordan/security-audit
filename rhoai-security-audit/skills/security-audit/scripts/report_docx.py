#!/usr/bin/env python3
"""Generate Word document security report using Red Hat template styling.

Produces a .docx file with Red Hat branding:
- Red Hat Display for headings
- Red Hat Text for body
- Red Hat Mono for code
- Severity-colored badges
- Finding tables with evidence and fix recommendations

Usage:
    python3 report_docx.py <scan-dir> -o security-report.docx
    python3 report_docx.py <scan-dir>  # outputs to <scan-dir>/security-report.docx
"""
import argparse
import json
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SEV_COLORS = {
    "critical": RGBColor(0xDC, 0x35, 0x45),
    "high": RGBColor(0xFD, 0x7E, 0x14),
    "medium": RGBColor(0xFF, 0xC1, 0x07),
    "low": RGBColor(0x17, 0xA2, 0xB8),
    "info": RGBColor(0x6C, 0x75, 0x7D),
}
SEV_RANK = {"critical": 4, "high": 3, "medium": 2, "low": 1, "info": 0}
RH_RED = RGBColor(0xEE, 0x00, 0x00)
RH_DARK = RGBColor(0x15, 0x15, 0x15)
RH_GREY = RGBColor(0x6A, 0x6E, 0x73)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEADING_FONT = "Red Hat Display"
BODY_FONT = "Red Hat Text"
MONO_FONT = "Red Hat Mono"


def load_findings(scan_dir):
    p = Path(scan_dir)
    triaged = p / "triaged-findings.json"
    if triaged.exists():
        return json.loads(triaged.read_text())
    for name in ["deduplicated-findings.json", "normalized-findings.json"]:
        f = p / name
        if f.exists():
            return json.loads(f.read_text())
    return []


def load_metadata(scan_dir):
    p = Path(scan_dir)
    meta = {}
    f = p / "scan-metadata.json"
    if f.exists():
        meta = json.loads(f.read_text())
    ss = p / "raw" / "security-summary.json"
    if ss.exists() and not meta.get("repo"):
        try:
            summary = json.loads(ss.read_text())
            meta.setdefault("repo", summary.get("repo", ""))
            meta.setdefault("date", summary.get("scan_date", ""))
            meta.setdefault("findings", summary.get("findings", {}))
        except Exception:
            pass
    ci = p / "raw" / "commit-info.json"
    if ci.exists():
        try:
            info = json.loads(ci.read_text())
            if not meta.get("branch"):
                meta["branch"] = info.get("default_branch", "main")
            if not meta.get("commit"):
                meta["commit"] = info.get("commit_sha", "")
        except Exception:
            pass
    if not meta.get("repo"):
        parts = Path(scan_dir).resolve().parts
        for i, part in enumerate(parts):
            if part == "output" and i + 1 < len(parts):
                meta["repo"] = f"opendatahub-io/{parts[i + 1]}"
                break
    if not meta.get("date"):
        parts = Path(scan_dir).resolve().parts
        for part in parts:
            if len(part) >= 10 and part[4] == "-" and part[7] == "-":
                meta["date"] = part[:10]
                break
    return meta


def _file_display(filepath, line_start):
    parts = filepath.replace("\\", "/").split("/")
    for i, p in enumerate(parts):
        if p in ("repo", "repos"):
            filepath = "/".join(parts[i + 2:]) if i + 2 < len(parts) else filepath
            break
    return f"{filepath}:{line_start}" if line_start else filepath


def _github_url(filepath, line_start, line_end, repo_full, ref):
    if not repo_full or not filepath:
        return ""
    url_path = filepath
    parts = filepath.replace("\\", "/").split("/")
    for i, p in enumerate(parts):
        if p in ("repo", "repos"):
            url_path = "/".join(parts[i + 2:]) if i + 2 < len(parts) else filepath
            break
    frag = f"#L{line_start}" if line_start else ""
    if line_end and line_end != line_start and line_start:
        frag = f"#L{line_start}-L{line_end}"
    return f"https://github.com/{repo_full}/blob/{ref}/{url_path}{frag}"


def _add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from lxml import etree
    part = paragraph.part

    # Split URL from fragment: Word/Pages encode # in Target, so
    # the fragment goes in w:anchor attribute instead
    base_url = url
    anchor = ""
    if "#" in url:
        base_url, anchor = url.split("#", 1)

    r_id = part.relate_to(base_url, RT.HYPERLINK, is_external=True)

    hyperlink = etree.SubElement(paragraph._element, qn("w:hyperlink"))
    hyperlink.set(qn("r:id"), r_id)
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)
    run_elem = etree.SubElement(hyperlink, qn("w:r"))
    rpr = etree.SubElement(run_elem, qn("w:rPr"))
    style = etree.SubElement(rpr, qn("w:rStyle"))
    style.set(qn("w:val"), "Hyperlink")
    color = etree.SubElement(rpr, qn("w:color"))
    color.set(qn("w:val"), "2563EB")
    u = etree.SubElement(rpr, qn("w:u"))
    u.set(qn("w:val"), "single")
    font = etree.SubElement(rpr, qn("w:rFonts"))
    font.set(qn("w:ascii"), MONO_FONT)
    font.set(qn("w:hAnsi"), MONO_FONT)
    sz = etree.SubElement(rpr, qn("w:sz"))
    sz.set(qn("w:val"), "18")
    t = etree.SubElement(run_elem, qn("w:t"))
    t.text = text
    t.set(qn("xml:space"), "preserve")


def _set_font(run, name=BODY_FONT, size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rpr.append(qn_elem("w:rFonts", {"w:ascii": name, "w:hAnsi": name, "w:cs": name}))


def qn_elem(tag, attribs):
    from lxml import etree
    elem = etree.SubElement(etree.Element("dummy"), qn(tag))
    for k, v in attribs.items():
        elem.set(qn(k), v)
    return elem


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = HEADING_FONT
        run.font.color.rgb = RH_DARK
    return h


def _add_body(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _add_code(doc, text):
    from lxml import etree
    lines = text.strip().split("\n")
    if len(lines) > 10:
        lines = lines[:10] + ["..."]

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)

    # Paragraph-level shading (grey background for entire block)
    ppr = p._element.get_or_add_pPr()
    shd = etree.SubElement(ppr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")

    # Paragraph borders (thin border around the code block)
    pbdr = etree.SubElement(ppr, qn("w:pBdr"))
    for side in ["top", "left", "bottom", "right"]:
        bdr = etree.SubElement(pbdr, qn(f"w:{side}"))
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "D0D0D0")

    run = p.add_run("\n".join(lines))
    run.font.name = MONO_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def _add_rich_text(doc, text):
    """Render text that may contain inline code blocks. Splits at Remediation: and indented lines."""
    if not text:
        return

    # Split at "Remediation:" boundary
    parts = text.split("Remediation:")
    prose = parts[0].strip()
    remediation = "Remediation:".join(parts[1:]).strip() if len(parts) > 1 else ""

    # Render prose (may still have indented code lines)
    if prose:
        _add_mixed_content(doc, prose)

    # Render remediation with label
    if remediation:
        _add_recommendation(doc, remediation)


def _add_mixed_content(doc, text):
    """Render text splitting prose from indented code blocks."""
    lines = text.split("\n")
    prose_buf = []
    code_buf = []

    for line in lines:
        is_code = (line.startswith("  ") and line.strip()) or line.startswith("\t")
        if is_code:
            if prose_buf:
                _add_body(doc, " ".join(prose_buf))
                prose_buf = []
            code_buf.append(line)
        else:
            if code_buf:
                _add_code(doc, "\n".join(code_buf))
                code_buf = []
            if line.strip():
                prose_buf.append(line.strip())

    if prose_buf:
        _add_body(doc, " ".join(prose_buf))
    if code_buf:
        _add_code(doc, "\n".join(code_buf))


def _add_recommendation(doc, text):
    """Render recommendation with a label, splitting prose from code."""
    p = doc.add_paragraph()
    run = p.add_run("Recommended Fix: ")
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RH_RED

    _add_mixed_content(doc, text)


def _set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    from lxml import etree
    shd = etree.SubElement(shading, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)


def _add_severity_table(doc, sev_counts, total):
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, sev in enumerate(["critical", "high", "medium", "low", "info"]):
        cell = table.rows[0].cells[i]
        count = sev_counts.get(sev, 0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{count}\n{sev.upper()}")
        run.font.name = HEADING_FONT
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = SEV_COLORS[sev]

    cell = table.rows[0].cells[5]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{total}\nTOTAL")
    run.font.name = HEADING_FONT
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RH_DARK

    return table


def _add_findings_table(doc, findings, repo_full, branch_ref, commit_ref):
    if not findings:
        _add_body(doc, "No findings in this category.")
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ["Severity", "Source", "File", "Title", "Description"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        _set_cell_shading(cell, "252525")

    sorted_findings = sorted(findings, key=lambda f: -SEV_RANK.get(f.get("severity", ""), 0))

    for f in sorted_findings[:50]:
        row = table.add_row()
        sev = f.get("severity", "info")
        triage = f.get("triage", {}).get("status", "")
        triage_label = {"corroborated": " [CORR]", "ai-only": " [AI]"}.get(triage, "")

        # Severity cell
        cell = row.cells[0]
        run = cell.paragraphs[0].add_run(sev.upper() + triage_label)
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = SEV_COLORS.get(sev, RH_GREY)

        # Source cell
        cell = row.cells[1]
        run = cell.paragraphs[0].add_run(f.get("source", ""))
        run.font.name = BODY_FONT
        run.font.size = Pt(9)

        # File cell (hyperlinked)
        cell = row.cells[2]
        display = _file_display(f.get("file", ""), f.get("line_start", 0))
        ref = branch_ref if f.get("origin") == "ai" else (commit_ref or branch_ref)
        url = _github_url(f.get("file", ""), f.get("line_start", 0),
                          f.get("line_end", 0), repo_full, ref)
        if url:
            _add_hyperlink(cell.paragraphs[0], url, display)
        else:
            run = cell.paragraphs[0].add_run(display)
            run.font.name = MONO_FONT
            run.font.size = Pt(8)

        # Title cell
        cell = row.cells[3]
        run = cell.paragraphs[0].add_run(f.get("title", "")[:60])
        run.font.name = BODY_FONT
        run.font.size = Pt(9)

        # Description cell
        cell = row.cells[4]
        desc = f.get("description", "")[:150]
        run = cell.paragraphs[0].add_run(desc)
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RH_GREY

    if len(findings) > 50:
        _add_body(doc, f"... and {len(findings) - 50} more findings not shown.", color=RH_GREY)


def _add_tool_coverage_table(doc, tool_sev):
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"

    headers = ["Tool", "Critical", "High", "Medium", "Low", "Info", "Total"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        _set_cell_shading(cell, "252525")

    for tool in sorted(tool_sev.keys()):
        s = tool_sev[tool]
        t = sum(s.values())
        row = table.add_row()
        run = row.cells[0].paragraphs[0].add_run(tool)
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.bold = True
        color = RH_GREY if t == 0 else RH_DARK
        run.font.color.rgb = color

        for j, sev in enumerate(["critical", "high", "medium", "low", "info"], 1):
            val = s.get(sev, 0)
            run = row.cells[j].paragraphs[0].add_run(str(val) if val else "")
            run.font.name = BODY_FONT
            run.font.size = Pt(9)
            if val:
                run.font.color.rgb = SEV_COLORS[sev]

        run = row.cells[6].paragraphs[0].add_run(str(t))
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.bold = True


def _add_finding_detail(doc, f, index, repo_full, branch_ref, commit_ref):
    """Render a single finding as a detailed section with metadata, description, snippet, and fix."""
    sev = f.get("severity", "")
    triage = f.get("triage", {}).get("status", "") if isinstance(f.get("triage"), dict) else ""
    triage_label = {"corroborated": " [CORR]", "ai-only": " [AI]"}.get(triage, "")
    title = f.get("title", f.get("id", ""))
    file_display = _file_display(f.get("file", ""), f.get("line_start", 0))

    _add_heading(doc, f"Finding {index}: {title} ({sev.upper()}{triage_label})", level=2)

    rule_id = f.get("rule_id", "")
    source = f.get("source", "")
    category = f.get("category", "")
    ref = branch_ref if f.get("origin") == "ai" else (commit_ref or branch_ref)
    url = _github_url(f.get("file", ""), f.get("line_start", 0),
                      f.get("line_end", 0), repo_full, ref)

    p = doc.add_paragraph()
    run = p.add_run("File: ")
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.bold = True
    if url:
        _add_hyperlink(p, url, file_display)
    else:
        run = p.add_run(file_display)
        run.font.name = MONO_FONT
        run.font.size = Pt(9)

    meta_parts = []
    if source:
        meta_parts.append(f"Source: {source}")
    if rule_id and rule_id != title:
        meta_parts.append(f"Rule: {rule_id}")
    if category:
        meta_parts.append(f"Category: {category}")
    if meta_parts:
        _add_body(doc, " | ".join(meta_parts), color=RH_GREY)

    desc = f.get("description", "")
    if not desc or desc == title:
        desc = f"{title} detected in {file_display} by {source}."
        if rule_id:
            desc += f" Rule: {rule_id}."
    _add_rich_text(doc, desc)

    snippet = f.get("snippet", "")
    if snippet:
        _add_code(doc, snippet[:500])

    rec = f.get("recommendation", "")
    if rec:
        _add_recommendation(doc, rec[:600])

    doc.add_paragraph()


def generate_docx(scan_dir, output_path, must_fix=False):
    findings = load_findings(scan_dir)
    metadata = load_metadata(scan_dir)

    repo_full = metadata.get("repo", "Unknown")
    repo_short = repo_full.split("/")[-1] if "/" in repo_full else repo_full
    branch_ref = metadata.get("branch", "main")
    commit_ref = metadata.get("commit", "")
    date = metadata.get("date", metadata.get("scan_date", ""))

    if must_fix:
        findings = [f for f in findings if SEV_RANK.get(f.get("severity", ""), 0) >= SEV_RANK["high"]]

    total = len(findings)
    sev_counts = Counter(f["severity"] for f in findings)
    triage_counts = Counter(
        f.get("triage", {}).get("status", "sast-only") if isinstance(f.get("triage"), dict) else "sast-only"
        for f in findings)

    doc = Document()

    # --- Cover Page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_text = "Must-Fix Security Report" if must_fix else "Security Audit Report"
    run = p.add_run(title_text)
    run.font.name = HEADING_FONT
    run.font.size = Pt(36)
    run.font.color.rgb = RH_DARK

    # Repo as hyperlink
    p = doc.add_paragraph()
    if repo_full and repo_full != "Unknown":
        _add_hyperlink(p, f"https://github.com/{repo_full}", repo_full)
    else:
        run = p.add_run(repo_full)
        run.font.name = HEADING_FONT
        run.font.size = Pt(18)
        run.font.color.rgb = RH_GREY

    # Branch + Commit + Date as labeled hyperlinks
    p = doc.add_paragraph()
    run = p.add_run("Branch: ")
    _set_font(run, BODY_FONT, 10, bold=True)
    if repo_full and repo_full != "Unknown":
        _add_hyperlink(p, f"https://github.com/{repo_full}/tree/{branch_ref}", branch_ref)
    else:
        run = p.add_run(branch_ref)
        _set_font(run, MONO_FONT, 10)

    run = p.add_run("  |  Commit: ")
    _set_font(run, BODY_FONT, 10, bold=True)
    commit_short = str(commit_ref)[:8]
    if repo_full and repo_full != "Unknown" and commit_ref:
        _add_hyperlink(p, f"https://github.com/{repo_full}/commit/{commit_ref}", commit_short)
    else:
        run = p.add_run(commit_short)
        _set_font(run, MONO_FONT, 10)

    run = p.add_run(f"  |  Date: {str(date)[:10]}")
    _set_font(run, BODY_FONT, 10)

    if must_fix:
        p = doc.add_paragraph()
        run = p.add_run("Scope: HIGH+ severity")
        _set_font(run, BODY_FONT, 10, bold=True, color=RH_RED)

    _add_body(doc, f"Generated by RHOAI Security Audit", color=RH_GREY)

    # Sensitivity banner
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("CONFIDENTIAL")
    _set_font(run, HEADING_FONT, 11, bold=True, color=SEV_COLORS["critical"])
    run = p.add_run(" — This report may contain undisclosed security findings. Do not share outside authorized personnel. Do not post in public channels.")
    _set_font(run, BODY_FONT, 10, color=RH_GREY)

    doc.add_page_break()

    # --- Executive Summary ---
    _add_heading(doc, "Executive Summary", level=1)
    _add_severity_table(doc, sev_counts, total)

    doc.add_paragraph()

    # Triage summary
    triage_parts = []
    if triage_counts.get("corroborated"):
        triage_parts.append(f"{triage_counts['corroborated']} corroborated (SAST + AI agree)")
    if triage_counts.get("ai-only"):
        triage_parts.append(f"{triage_counts['ai-only']} AI-only (logic bugs SAST missed)")
    sast_only = triage_counts.get("sast-only", 0)
    if sast_only:
        triage_parts.append(f"{sast_only} SAST-only")
    demoted = sum(1 for f in findings if isinstance(f.get("triage"), dict) and f["triage"].get("demoted_from"))
    if demoted:
        triage_parts.append(f"{demoted} demoted (non-production code)")

    if triage_parts:
        _add_heading(doc, "Triage Summary", level=2)
        for part in triage_parts:
            _add_body(doc, f"  {part}")

    # --- Detailed Findings ---
    sorted_findings = sorted(findings, key=lambda f: -SEV_RANK.get(f.get("severity", ""), 0))

    if must_fix:
        _add_heading(doc, f"Must-Fix Findings ({len(sorted_findings)})", level=1)
        _add_body(doc, "These findings require immediate attention.")
        doc.add_paragraph()

        for i, f in enumerate(sorted_findings, 1):
            _add_finding_detail(doc, f, i, repo_full, branch_ref, commit_ref)
    else:
        # Critical + High with full detail
        crit_high = [f for f in sorted_findings if f["severity"] in ("critical", "high")]
        if crit_high:
            _add_heading(doc, f"Critical and High Findings ({len(crit_high)})", level=1)
            _add_body(doc, "These findings require immediate attention.")
            doc.add_paragraph()

            for i, f in enumerate(crit_high, 1):
                _add_finding_detail(doc, f, i, repo_full, branch_ref, commit_ref)

        # AI Review Findings (medium/low)
        ai_findings = [f for f in findings if f.get("origin") == "ai" and f["severity"] not in ("critical", "high")]
        if ai_findings:
            _add_heading(doc, f"AI Review Findings ({len(ai_findings)})", level=1)
            _add_body(doc, "Findings from adversarial multi-agent review and semantic security analysis.")
            doc.add_paragraph()
            _add_findings_table(doc, ai_findings, repo_full, branch_ref, commit_ref)

        # Dependency Vulnerabilities
        sca_findings = [f for f in findings if f.get("category") == "sca"]
        if sca_findings:
            _add_heading(doc, f"Dependency Vulnerabilities ({len(sca_findings)} CVEs)", level=1)
            _add_findings_table(doc, sca_findings, repo_full, branch_ref, commit_ref)

        # All Other Findings
        other = [f for f in findings
                 if f["severity"] not in ("critical", "high")
                 and f.get("origin") != "ai"
                 and f.get("category") != "sca"]
        if other:
            _add_heading(doc, f"Other Findings ({len(other)})", level=1)
            _add_findings_table(doc, other, repo_full, branch_ref, commit_ref)

    # --- Tool Coverage ---
    _add_heading(doc, "Tool Coverage", level=1)
    tool_sev = defaultdict(Counter)
    for f in findings:
        tool_sev[f.get("source", "unknown")][f["severity"]] += 1
    meta_findings = metadata.get("findings", {})
    for tk, tc in meta_findings.items():
        tn = tk.replace("_", "-")
        if tn not in tool_sev:
            tool_sev[tn] = Counter()
    _add_tool_coverage_table(doc, tool_sev)

    # Save
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Generate Word document security report")
    parser.add_argument("scan_dir", help="Scan output directory")
    parser.add_argument("-o", "--output", help="Output file path")
    parser.add_argument("--must-fix", action="store_true", help="Generate must-fix report (HIGH+ severity only)")
    args = parser.parse_args()

    output = args.output
    if not output:
        filename = "must-fix-report.docx" if args.must_fix else "security-report.docx"
        output = str(Path(args.scan_dir) / filename)

    path = generate_docx(args.scan_dir, output, must_fix=args.must_fix)
    print(f"Report saved to: {path}")


if __name__ == "__main__":
    main()
